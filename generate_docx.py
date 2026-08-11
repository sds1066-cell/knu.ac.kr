from __future__ import annotations

import hashlib
import re
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt
from lxml import etree

OUTPUT_DIR = Path("output")
OUTPUT_PATH = OUTPUT_DIR / "第二章_2.2_韩语忠实翻译稿_重新输出.docx"
TEMP_PATH = OUTPUT_DIR / "_base.docx"

BODY_FONT = "Batang"
BODY_SIZE = 10.5
FOOTNOTE_SIZE = 9.0
W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
CT_NS = "http://schemas.openxmlformats.org/package/2006/content-types"
PKG_REL_NS = "http://schemas.openxmlformats.org/package/2006/relationships"
NS = {"w": W_NS, "r": R_NS}


def set_run_font(run, size: float = BODY_SIZE, *, bold: bool | None = None, italic: bool | None = None) -> None:
    run.font.name = BODY_FONT
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{attr}"), BODY_FONT)


def set_style_font(style, size: float) -> None:
    style.font.name = BODY_FONT
    style.font.size = Pt(size)
    rpr = style._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{attr}"), BODY_FONT)


def add_footnote_reference(paragraph, footnote_number: int, internal_id: int) -> None:
    run = paragraph.add_run()
    set_run_font(run)
    rpr = run._element.get_or_add_rPr()
    rstyle = OxmlElement("w:rStyle")
    rstyle.set(qn("w:val"), "FootnoteReference")
    rpr.append(rstyle)
    ref = OxmlElement("w:footnoteReference")
    ref.set(qn("w:id"), str(internal_id))
    run._element.append(ref)


def add_segments(paragraph, segments: list[tuple[str, str | int]], id_map: dict[int, int]) -> None:
    for kind, value in segments:
        if kind == "text":
            run = paragraph.add_run(str(value))
            set_run_font(run)
        elif kind == "fn":
            number = int(value)
            add_footnote_reference(paragraph, number, id_map[number])
        else:
            raise ValueError(f"Unknown segment kind: {kind}")


def make_document() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(25)
    section.bottom_margin = Mm(25)
    section.left_margin = Mm(30)
    section.right_margin = Mm(25)
    section.header_distance = Mm(12.7)
    section.footer_distance = Mm(12.7)

    normal = doc.styles["Normal"]
    set_style_font(normal, BODY_SIZE)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.line_spacing = 1.6
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.first_line_indent = Pt(BODY_SIZE * 2)

    try:
        footnote_text = doc.styles["Footnote Text"]
    except KeyError:
        footnote_text = doc.styles.add_style("Footnote Text", WD_STYLE_TYPE.PARAGRAPH)
    set_style_font(footnote_text, FOOTNOTE_SIZE)
    footnote_text.paragraph_format.line_spacing = 1.0
    footnote_text.paragraph_format.space_before = Pt(0)
    footnote_text.paragraph_format.space_after = Pt(0)
    footnote_text.paragraph_format.first_line_indent = Pt(0)

    try:
        footnote_ref = doc.styles["Footnote Reference"]
    except KeyError:
        footnote_ref = doc.styles.add_style("Footnote Reference", WD_STYLE_TYPE.CHARACTER)
    set_style_font(footnote_ref, FOOTNOTE_SIZE)
    footnote_ref.font.superscript = True

    doc.core_properties.title = "제2장 2.2 관련 판례의 자제 의무 해석 기능"
    doc.core_properties.subject = "중국어 원문에 충실한 한국어 번역문"
    doc.core_properties.language = "ko-KR"

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.paragraph_format.first_line_indent = Pt(0)
    title.paragraph_format.line_spacing = 1.3
    title.paragraph_format.space_after = Pt(10)
    title_run = title.add_run("2.2 관련 판례의 자제 의무 해석 기능")
    set_run_font(title_run, 12.0, bold=True)

    footnote_numbers = list(range(13, 26))
    id_map = {number: idx + 1 for idx, number in enumerate(footnote_numbers)}

    paragraphs: list[list[tuple[str, str | int]]] = [
        [
            ("text", "제74조 제3항과 제83조 제3항은 관련국이 실질적인 잠정약정을 체결하기 위해 노력하는 동시에 최종 합의를 위태롭게 하거나 방해하지 않을 것을 요구한다."),
            ("fn", 13),
            ("text", " 조약은 구체적인 위반 유형을 열거하지 않으므로, 사법 및 중재 실무는 행위의 경계를 명확히 하는 역할을 담당해 왔다. 활동의 명칭만을 기준으로 판단할 경우, 자제의무를 현실적 활동의 전면적인 동결로 해석할 수도 있고, 활동의 물리적 결과와 국가의 대응이 협상 조건에 미치는 영향을 간과할 수도 있다. 관련 판례는 이러한 문제를 중심으로 점차 보다 명확한 판단 경로를 형성해 왔다."),
        ],
        [
            ("text", "절차 단계에 따라 재판기관이 다루어야 할 쟁점은 달라진다. 1976년 「에게해 대륙붕 사건」은 권리의 보전과 회복 불가능한 손해를 심사하였을 뿐, 제74조 제3항 또는 제83조 제3항에 따른 실체적 책임은 다루지 않았다. 재판소는 문제된 탄성파 탐사가 일시적 성격을 띠고, 해저에 시설을 설치하지 않았으며, 천연자원을 실제로 점유하거나 이용하지 않았다는 점에 주목하였다. 또한 당시 자료 역시 해당 탐사가 해저, 하층토 또는 천연자원에 물리적 손상을 초래할 것임을 보여 주지 않았다. 그리스가 주장한 권리 침해가 성립하더라도 적절한 방식으로 구제될 수 있었으므로, 재판소는 잠정조치를 명하지 않았다. 재판소는 동시에 신청 범위의 제약을 받아 그리스가 제기한 무력 문제에 관하여 실체적 판단을 내리지 않았다."),
            ("fn", 14),
            ("text", " 1978년 같은 사건의 판결은 재판소의 관할권만을 다루었다."),
            ("fn", 15),
            ("text", " 이 두 재판은 보전 단계에서 일시적 활동이 어떻게 평가되는지를 보여 주지만, 자제의무의 실체적 기준을 대신할 수 없으며, 탄성파 탐사가 일반적인 의미에서 적법하다는 결론을 뒷받침하지도 않는다. 일방적 활동 자체의 한계를 판단하려면 활동의 효과와 국가의 대응 방식을 직접 평가한 사건으로 나아가야 한다."),
        ],
        [
            ("text", "「가이아나-수리남 사건」은 쟁점을 실체적 평가의 단계로 진전시켰다. 중재재판소는 관련국이 잠정약정 체결을 위해 선의의 협의를 진행하였는지, 일방적 활동이 최종 합의를 위태롭게 하거나 방해하였는지, 그리고 퇴거 행위가 무력의 위협 또는 사용을 구성하였는지를 각각 심사하였다. 잠정약정 체결을 위해 노력할 의무는 협의가 선의, 타협적 태도 및 필요한 양보를 바탕으로 진행되었는지를 살피며, 잠정약정을 체결하지 못했다는 사실만으로는 위반으로 인정하기에 부족하다. 반면 자제의무는 일방적 활동이 협상 조건과 최종 합의의 형성 여지를 변경하였는지에 초점을 둔다."),
            ("fn", 16),
            ("text", " 수리남이 CGX 시추 플랫폼을 퇴거시킨 행위는 별도로 「유엔헌장」상 무력의 위협 또는 사용에 관한 규범에 따라 심사되었다."),
            ("fn", 17),
            ("text", " 세 가지 쟁점은 동일한 사건에서 비롯되었지만, 적용 근거와 입증 경로는 서로 다르다. 활동의 효과에 따라 자제의무가 평가되며, 퇴거 또는 법집행에 의한 저지는 관련 국제법 규범에 따라 별도로 판단되어야 한다. 대응이 사실상 상대방의 활동을 배제하였는지는 협상 환경의 변화를 살펴보는 사실적 배경이 될 수 있으나, 이를 자제의무와 결합하여 하나의 법적 결론으로 볼 수는 없다."),
        ],
        [
            ("text", "중재재판소는 경계미획정 해역을 반드시 정지 상태로 유지되어야 하는 공간으로 이해하지 않았다. 판정은 해양환경에 물리적 변화를 초래하지 않는 일방적 행위는 일반적으로 최종 합의를 방해하지 않지만, 물리적 변화를 초래할 수 있는 행위는 현상을 변경하고 상대방의 협상 지위를 약화시킬 수 있으므로 통상 양측이 공동으로 수행하거나 합의에 따라 수행해야 한다고 보았다. 이에 따라 중재재판소는 탄성파 탐사와 시험굴착을 구분하여, 전자는 자제의무와 양립하지 않는 것으로 인정되지 않았으나 후자는 영구적인 변화 또는 손상을 초래할 수 있다고 보았다."),
            ("fn", 18),
            ("text", " 이러한 차이의 핵심은 사실상태를 쉽게 회복할 수 있는지 여부에 있다. 되돌리기 어려운 물리적 변화는 원상회복 비용을 높이고, 본래 협상을 통해 조정할 수 있었던 상태를 더욱 변경하기 어렵게 만든다. 일시성, 가역성 및 보상 가능성은 회복과 조정을 위한 여지를 남기지만, 그 자체로 활동의 적법 여부를 자동으로 결정하지는 않는다. 본 논문은 이에 근거하여 지속적·배타적 또는 통제적 효과를 후속 사실 비교에서 확인해야 할 문제로 설정한다. 구체적인 판단은 여전히 공간적 위치, 기술적 속성, 지속기간 및 실제 영향과 결합하여 이루어져야 한다. Nishimoto 역시 영구성을 중요하지만 배타적이지 않은 요소로 보며, 절차, 청구 내용 및 사건의 맥락이 모두 결론에 영향을 미친다고 지적한다."),
            ("fn", 19),
        ],
        [
            ("text", "문제된 활동 자체에 위험성이 있더라도, 국가는 이를 이유로 저지 방식을 임의로 선택할 수는 없다. 「가이아나-수리남 사건」에서 수리남 군함은 CGX 플랫폼에 정해진 시간 안에 해당 해역을 떠날 것을 요구하고, 이에 불응할 경우 그 결과를 감수해야 할 것이라고 위협하였다. 중재재판소는 명령을 내린 주체의 권한 근거, 명령의 내용, 위협의 임박성 및 그로 인해 발생할 수 있는 결과를 검토한 후, 해당 행위가 군사행동의 위협을 구성하며 일반적인 해상 법집행에는 해당하지 않는다고 판단하였다."),
            ("fn", 20),
            ("text", " 중재재판소는 또한 당시 협상, 「협약」 제15부 및 부속서 Ⅶ에 따른 분쟁해결절차의 개시 또는 잠정조치 신청과 같은 평화적 수단으로 분쟁을 처리할 수 있었다고 지적하였다."),
            ("fn", 21),
            ("text", " 관련 연구 역시 활동의 물리적 결과와 퇴거·법집행 대응을 서로 인접하지만 상호 대체할 수 없는 문제로 본다."),
            ("fn", 22),
            ("text", " 강제적 대응을 평가할 때에는 행위 주체의 권한, 조치의 강도, 임박성 및 이용 가능한 평화적 대체수단을 별도로 확인해야 한다. 현장 활동이 저지되었다는 결과만으로는 통제 관계 또는 관할권이 이미 변경되었음을 입증할 수 없다."),
        ],
        [
            ("text", "「가나-코트디부아르 해양경계획정 사건」의 두 단계는 영구성이 서로 다른 절차에서 상이한 기능을 수행한다는 점을 더욱 분명히 보여 준다. 2015년 잠정조치 명령은 새로운 굴착이 초래할 수 있는 중대하고 영구적인 물리적 변화를 권리 보전 단계의 핵심 위험으로 보았으며, 가나에 분쟁수역에서 더 이상의 새로운 굴착이 이루어지지 않도록 보장할 것을 요구하였다. 특별재판부는 이미 굴착이 진행된 사업을 갑자기 중단할 경우 중대한 경제적 손실과 해양환경 위험이 발생할 수 있다는 이유로 기존 활동을 전면 중단시키지는 않았다."),
            ("fn", 23),
            ("text", " 2017년 본안판결은 제83조 제3항에 포함된 두 의무를 각각 심사하였다. 코트디부아르는 가나에 잠정약정에 관한 협상을 요청하지 않았으며, 가나는 2015년 명령에 따라 새로운 굴착을 이미 중단하였다. 최종 경계획정 결과에 따르면 문제된 석유·가스 활동이 이루어진 해역은 가나에 귀속되었고, 코트디부아르의 최종 청구에서 지칭한 “코트디부아르 해역”에는 포함되지 않았다. 이에 특별재판부는 자제의무 위반을 인정하지 않았다."),
            ("fn", 24),
            ("text", " 동일한 요소는 잠정조치 단계에서 위험을 식별하는 데 사용되지만, 본안판결 단계에서는 청구 범위, 활동이 이루어진 공간 및 국가의 후속 행위와 결합되어야만 책임의 근거가 될 수 있다."),
        ],
        [
            ("text", "책임에 관한 결론은 사실의 입증을 거쳐야 한다. 「소말리아-케냐 사건」은 굴착이 영구적인 물리적 변화를 초래할 수 있음을 인정하는 한편, 관련 활동의 시점, 위치 및 효과가 입증되어야 한다고 보았다. 이 사건에서 확인된 광구에 대한 양허 부여, 탄성파 탐사 및 기타 조사활동은 최종 합의에 충분한 영향을 미쳤다는 점이 입증되지 않았으며, 굴착과 관련된 구체적인 시점과 위치도 충분히 입증되지 않았다."),
            ("fn", 25),
            ("text", " 전자의 주장은 효과에 관한 증거가 부족했고, 후자의 주장은 시간적·공간적 증거가 부족했다. “충분히 입증되지 않았다”는 것은 책임에 관한 결론을 제한할 뿐, 관련 사실이 전혀 발생하지 않았다는 뜻은 아니다. 활동 유형은 일반적인 위험을 시사할 뿐이며, 행위의 발생, 공간적 귀속 및 실제 효과가 입증되어야만 그 위험을 특정 국가에 귀속시킬 수 있다."),
        ],
        [
            ("text", "이들 판례는 공통적으로 하나의 기본 판단을 제시한다. 자제의무가 보호하는 것은 최종 합의를 계속 협상하고 조정할 수 있는 여지이므로, 분쟁 상태를 회복하기 어렵게 만들 수 있는 행위가 그 관심의 대상이 된다. 물리적 변화의 회복 가능성, 절차 단계 및 증거 상황이 함께 평가의 강도를 결정하며, 강제적 대응은 별도의 규범적 심사를 받아야 한다. 고정시설, 조사, 통항 및 접근 제한이 서로 비교 가능한지는 구체적인 활동, 대응 방식 및 사실적 근거를 결합하여 각각 판단해야 한다. 판례는 행위가 어떻게 평가되는지를 설명할 수 있지만, 한중 양국이 왜 우선 어업 분야에서 협력을 형성했는지는 설명할 수 없다. 이 문제에 답하기 위해서는 황해의 현실적 조건과 양국의 협상과정을 함께 살펴보고, 어업 의제가 왜 최종 경계획정에 앞서 협력약정에 포함될 수 있었는지를 검토해야 한다."),
        ],
    ]

    for segments in paragraphs:
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.first_line_indent = Pt(BODY_SIZE * 2)
        paragraph.paragraph_format.line_spacing = 1.6
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        add_segments(paragraph, segments, id_map)

    doc.save(TEMP_PATH)
    inject_footnotes(TEMP_PATH, OUTPUT_PATH, id_map)
    TEMP_PATH.unlink(missing_ok=True)


def add_xml_run(parent, text: str, *, italic: bool = False, style: str | None = None, size_half_points: int = 18) -> None:
    run = etree.SubElement(parent, f"{{{W_NS}}}r")
    rpr = etree.SubElement(run, f"{{{W_NS}}}rPr")
    rfonts = etree.SubElement(rpr, f"{{{W_NS}}}rFonts")
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(f"{{{W_NS}}}{attr}", BODY_FONT)
    sz = etree.SubElement(rpr, f"{{{W_NS}}}sz")
    sz.set(f"{{{W_NS}}}val", str(size_half_points))
    szcs = etree.SubElement(rpr, f"{{{W_NS}}}szCs")
    szcs.set(f"{{{W_NS}}}val", str(size_half_points))
    if italic:
        etree.SubElement(rpr, f"{{{W_NS}}}i")
        etree.SubElement(rpr, f"{{{W_NS}}}iCs")
    if style:
        rstyle = etree.SubElement(rpr, f"{{{W_NS}}}rStyle")
        rstyle.set(f"{{{W_NS}}}val", style)
    t = etree.SubElement(run, f"{{{W_NS}}}t")
    if text.startswith(" ") or text.endswith(" "):
        t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    t.text = text


def append_markdown_runs(parent, text: str) -> None:
    position = 0
    for match in re.finditer(r"\*([^*]+)\*", text):
        if match.start() > position:
            add_xml_run(parent, text[position:match.start()])
        add_xml_run(parent, match.group(1), italic=True)
        position = match.end()
    if position < len(text):
        add_xml_run(parent, text[position:])


def build_footnotes_xml(id_map: dict[int, int]) -> bytes:
    footnote_texts: dict[int, str] = {
        13: "United Nations, *United Nations Convention on the Law of the Sea*, 10 December 1982, Arts. 74(3), 83(3), official English text, pp. 52, 56; 联合国, 『联合国海洋法公约』, 제74조 제3항 및 제83조 제3항, 유엔 중국어 공식본 PDF 55, 59쪽. 관련 규범 문구는 영어 공식본을 기준으로 하며, 중국어 용어는 유엔 중국어 공식본을 참조하였다.",
        14: "*Aegean Sea Continental Shelf (Greece v. Turkey)*, Interim Protection, Order of 11 September 1976, I.C.J. Reports 1976, paras. 22–35, especially paras. 30, 33–35.",
        15: "*Aegean Sea Continental Shelf (Greece v. Turkey)*, Judgment of 19 December 1978, I.C.J. Reports 1978, opening title and “Jurisdiction of the Court.”",
        16: "*Guyana v. Suriname*, Award of the Arbitral Tribunal, 17 September 2007, paras. 459–466.",
        17: "*Guyana v. Suriname*, Award of the Arbitral Tribunal, 17 September 2007, paras. 433–445, especially paras. 439, 445.",
        18: "*Guyana v. Suriname*, Award of the Arbitral Tribunal, 17 September 2007, paras. 465–470, 479–481.",
        19: "Kentaro Nishimoto, “The Obligation of Self-Restraint in Undelimited Maritime Areas,” *Japan Review*, Vol. 3, No. 1, 2019, pp. 33–36, especially pp. 33–35.",
        20: "*Guyana v. Suriname*, Award of the Arbitral Tribunal, 17 September 2007, paras. 433–445, especially paras. 439, 445.",
        21: "*Guyana v. Suriname*, Award of the Arbitral Tribunal, 17 September 2007, paras. 482–484.",
        22: "김민철, 「경계미획정 수역에 관한 자제의무 해석론의 변천과 현주소: ICJ의 2021년 소말리아-케냐 사건 본안판결을 중심으로」, 『서울국제법연구』 제30권 제1호 (2023), 85–88쪽, 특히 86쪽.",
        23: "*Dispute Concerning Delimitation of the Maritime Boundary between Ghana and Côte d’Ivoire in the Atlantic Ocean (Ghana/Côte d’Ivoire)*, Provisional Measures, Order of 25 April 2015, ITLOS Reports 2015, paras. 88–104, 108(1)(a).",
        24: "*Dispute Concerning Delimitation of the Maritime Boundary between Ghana and Côte d’Ivoire in the Atlantic Ocean (Ghana/Côte d’Ivoire)*, Judgment of 23 September 2017, ITLOS Reports 2017, paras. 624–634, especially paras. 626–634.",
        25: "*Maritime Delimitation in the Indian Ocean (Somalia v. Kenya)*, Judgment of 12 October 2021, I.C.J. Reports 2021, paras. 198–211, especially paras. 205–211.",
    }

    root = etree.Element(f"{{{W_NS}}}footnotes", nsmap={"w": W_NS})

    separator = etree.SubElement(root, f"{{{W_NS}}}footnote")
    separator.set(f"{{{W_NS}}}id", "-1")
    separator.set(f"{{{W_NS}}}type", "separator")
    p = etree.SubElement(separator, f"{{{W_NS}}}p")
    r = etree.SubElement(p, f"{{{W_NS}}}r")
    etree.SubElement(r, f"{{{W_NS}}}separator")

    continuation = etree.SubElement(root, f"{{{W_NS}}}footnote")
    continuation.set(f"{{{W_NS}}}id", "0")
    continuation.set(f"{{{W_NS}}}type", "continuationSeparator")
    p = etree.SubElement(continuation, f"{{{W_NS}}}p")
    r = etree.SubElement(p, f"{{{W_NS}}}r")
    etree.SubElement(r, f"{{{W_NS}}}continuationSeparator")

    for number in range(13, 26):
        footnote = etree.SubElement(root, f"{{{W_NS}}}footnote")
        footnote.set(f"{{{W_NS}}}id", str(id_map[number]))
        p = etree.SubElement(footnote, f"{{{W_NS}}}p")
        ppr = etree.SubElement(p, f"{{{W_NS}}}pPr")
        pstyle = etree.SubElement(ppr, f"{{{W_NS}}}pStyle")
        pstyle.set(f"{{{W_NS}}}val", "FootnoteText")
        spacing = etree.SubElement(ppr, f"{{{W_NS}}}spacing")
        spacing.set(f"{{{W_NS}}}before", "0")
        spacing.set(f"{{{W_NS}}}after", "0")
        spacing.set(f"{{{W_NS}}}line", "240")
        spacing.set(f"{{{W_NS}}}lineRule", "auto")

        ref_run = etree.SubElement(p, f"{{{W_NS}}}r")
        ref_rpr = etree.SubElement(ref_run, f"{{{W_NS}}}rPr")
        ref_style = etree.SubElement(ref_rpr, f"{{{W_NS}}}rStyle")
        ref_style.set(f"{{{W_NS}}}val", "FootnoteReference")
        etree.SubElement(ref_run, f"{{{W_NS}}}footnoteRef")

        tab_run = etree.SubElement(p, f"{{{W_NS}}}r")
        etree.SubElement(tab_run, f"{{{W_NS}}}tab")
        append_markdown_runs(p, footnote_texts[number])

    return etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone="yes")


def inject_footnotes(source: Path, destination: Path, id_map: dict[int, int]) -> None:
    with zipfile.ZipFile(source, "r") as zin:
        files = {name: zin.read(name) for name in zin.namelist()}

    content_types = etree.fromstring(files["[Content_Types].xml"])
    if not content_types.xpath("ct:Override[@PartName='/word/footnotes.xml']", namespaces={"ct": CT_NS}):
        override = etree.SubElement(content_types, f"{{{CT_NS}}}Override")
        override.set("PartName", "/word/footnotes.xml")
        override.set("ContentType", "application/vnd.openxmlformats-officedocument.wordprocessingml.footnotes+xml")
    files["[Content_Types].xml"] = etree.tostring(content_types, xml_declaration=True, encoding="UTF-8", standalone="yes")

    rels_path = "word/_rels/document.xml.rels"
    rels = etree.fromstring(files[rels_path])
    ids = []
    for rel in rels.findall(f"{{{PKG_REL_NS}}}Relationship"):
        match = re.fullmatch(r"rId(\d+)", rel.get("Id", ""))
        if match:
            ids.append(int(match.group(1)))
    next_id = max(ids, default=0) + 1
    relationship = etree.SubElement(rels, f"{{{PKG_REL_NS}}}Relationship")
    relationship.set("Id", f"rId{next_id}")
    relationship.set("Type", "http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes")
    relationship.set("Target", "footnotes.xml")
    files[rels_path] = etree.tostring(rels, xml_declaration=True, encoding="UTF-8", standalone="yes")

    settings = etree.fromstring(files["word/settings.xml"])
    existing = settings.find(f"{{{W_NS}}}footnotePr")
    if existing is not None:
        settings.remove(existing)
    footnote_pr = etree.SubElement(settings, f"{{{W_NS}}}footnotePr")
    num_fmt = etree.SubElement(footnote_pr, f"{{{W_NS}}}numFmt")
    num_fmt.set(f"{{{W_NS}}}val", "decimal")
    num_start = etree.SubElement(footnote_pr, f"{{{W_NS}}}numStart")
    num_start.set(f"{{{W_NS}}}val", "13")
    num_restart = etree.SubElement(footnote_pr, f"{{{W_NS}}}numRestart")
    num_restart.set(f"{{{W_NS}}}val", "continuous")
    files["word/settings.xml"] = etree.tostring(settings, xml_declaration=True, encoding="UTF-8", standalone="yes")

    files["word/footnotes.xml"] = build_footnotes_xml(id_map)

    with zipfile.ZipFile(destination, "w", compression=zipfile.ZIP_DEFLATED) as zout:
        for name, data in files.items():
            zout.writestr(name, data)


def validate() -> None:
    with zipfile.ZipFile(OUTPUT_PATH, "r") as zf:
        bad = zf.testzip()
        if bad is not None:
            raise RuntimeError(f"Corrupt ZIP member: {bad}")
        names = set(zf.namelist())
        required = {
            "[Content_Types].xml",
            "word/document.xml",
            "word/styles.xml",
            "word/settings.xml",
            "word/footnotes.xml",
            "word/_rels/document.xml.rels",
        }
        missing = required - names
        if missing:
            raise RuntimeError(f"Missing required DOCX parts: {sorted(missing)}")
        document_xml = etree.fromstring(zf.read("word/document.xml"))
        footnotes_xml = etree.fromstring(zf.read("word/footnotes.xml"))
        refs = document_xml.xpath("//w:footnoteReference", namespaces=NS)
        notes = footnotes_xml.xpath("//w:footnote[not(@w:type)]", namespaces=NS)
        if len(refs) != 13:
            raise RuntimeError(f"Expected 13 footnote references, found {len(refs)}")
        if len(notes) != 13:
            raise RuntimeError(f"Expected 13 footnotes, found {len(notes)}")

    reopened = Document(OUTPUT_PATH)
    visible_paragraphs = [p for p in reopened.paragraphs if p.text.strip()]
    if len(visible_paragraphs) != 9:
        raise RuntimeError(f"Expected 1 heading + 8 body paragraphs, found {len(visible_paragraphs)}")
    if visible_paragraphs[0].text.strip() != "2.2 관련 판례의 자제 의무 해석 기능":
        raise RuntimeError("Heading mismatch")
    digest = hashlib.sha256(OUTPUT_PATH.read_bytes()).hexdigest()
    print(f"OUTPUT={OUTPUT_PATH}")
    print(f"SIZE_BYTES={OUTPUT_PATH.stat().st_size}")
    print(f"SHA256={digest}")
    print("DOCX_VALIDATION=PASS")


if __name__ == "__main__":
    make_document()
    validate()
